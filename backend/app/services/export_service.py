import io
import re
import uuid
from typing import Literal
from bs4 import BeautifulSoup, NavigableString, Tag

# ── Unit helpers ───────────────────────────────────────────────────────────────

def _to_cm(value: float, unit: str) -> float:
    if unit == "mm": return value / 10
    if unit == "in": return value * 2.54
    return value  # already cm

def _to_pt(value: float, unit: str = "px") -> float:
    """Convert to points (1pt = 1.333px at 96dpi)."""
    if unit == "px": return value * 0.75
    if unit == "pt": return value
    if unit == "em": return value * 12  # assume 12pt base
    return value

def _parse_style(style_str: str) -> dict:
    result = {}
    for part in (style_str or "").split(";"):
        part = part.strip()
        if ":" in part:
            k, v = part.split(":", 1)
            result[k.strip().lower()] = v.strip()
    return result

def _parse_pt(value_str: str) -> float:
    """Parse a CSS size string and return pt value."""
    value_str = value_str.strip()
    try:
        if value_str.endswith("pt"):
            return float(value_str[:-2])
        if value_str.endswith("px"):
            return float(value_str[:-2]) * 0.75
        if value_str.endswith("em"):
            return float(value_str[:-2]) * 12
        if value_str.endswith("cm"):
            return float(value_str[:-2]) * 28.35
        if value_str.endswith("mm"):
            return float(value_str[:-2]) * 2.835
        if value_str.endswith("in"):
            return float(value_str[:-2]) * 72
        return float(value_str) * 0.75
    except (ValueError, AttributeError):
        return 12.0

def _parse_cm(value_str: str) -> float:
    """Parse a CSS size string and return cm value."""
    value_str = (value_str or "").strip()
    try:
        if value_str.endswith("cm"): return float(value_str[:-2])
        if value_str.endswith("mm"): return float(value_str[:-2]) / 10
        if value_str.endswith("in"): return float(value_str[:-2]) * 2.54
        if value_str.endswith("px"): return float(value_str[:-2]) / 37.795
        if value_str.endswith("pt"): return float(value_str[:-2]) / 28.35
        if value_str.endswith("em"): return float(value_str[:-2]) * 0.423
        return 0.0
    except (ValueError, AttributeError):
        return 0.0

# ── EPUB HTML sanitiser ────────────────────────────────────────────────────────

def _sanitize_html_to_xhtml(raw_html: str) -> str:
    """
    Re-serialise lenient editor HTML (TipTap/HTML5) as XHTML-safe markup.

    Uses BeautifulSoup's tolerant html.parser to fix unclosed tags, bare text,
    etc., then emits self-closed void elements and properly escaped text so the
    result can be embedded inside a strict XHTML document without lxml choking.

    Always returns at least one element — lxml raises "Document is empty" when
    the body is truly empty, so we substitute a non-breaking space paragraph.
    """
    import html as _h

    VOID = frozenset({
        "area", "base", "br", "col", "embed", "hr", "img",
        "input", "link", "meta", "param", "source", "track", "wbr",
    })

    if not raw_html or not raw_html.strip():
        return "<p>&#160;</p>"

    # Strip page-break spacer divs inserted by the editor
    raw_html = re.sub(
        r'<div[^>]*data-type="page-break-spacer"[^>]*>.*?</div>',
        "", raw_html, flags=re.DOTALL,
    ).strip()

    if not raw_html:
        return "<p>&#160;</p>"

    soup = BeautifulSoup(raw_html, "html.parser")

    def _emit(node) -> str:
        if isinstance(node, NavigableString):
            return _h.escape(str(node))
        if not isinstance(node, Tag):
            return ""
        tag = node.name
        if not tag:
            return ""
        attrs = ""
        for k, v in (node.attrs or {}).items():
            if isinstance(v, list):
                v = " ".join(v)
            attrs += f' {k}="{_h.escape(str(v))}"'
        if tag in VOID:
            return f"<{tag}{attrs}/>"
        inner = "".join(_emit(c) for c in node.children)
        return f"<{tag}{attrs}>{inner}</{tag}>"

    parts = []
    for child in soup.children:
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                parts.append(f"<p>{_h.escape(text)}</p>")
        elif isinstance(child, Tag):
            xhtml = _emit(child)
            if xhtml.strip():
                parts.append(xhtml)

    result = "\n".join(parts)
    return result if result.strip() else "<p>&#160;</p>"


# ── Header / footer helpers ───────────────────────────────────────────────────

def _to_roman(n: int) -> str:
    vals = [1000,900,500,400,100,90,50,40,10,9,5,4,1]
    syms = ["m","cm","d","cd","c","xc","l","xl","x","ix","v","iv","i"]
    result = ""
    for v, s in zip(vals, syms):
        while n >= v:
            result += s
            n -= v
    return result


def _format_page_number(page: int, total: int, fmt: str) -> str:
    if fmt == "number-of-total":
        return f"{page} of {total}"
    if fmt == "roman":
        return _to_roman(max(page, 1))
    return str(page)


def _resolve_hf_tokens(text: str, doc_title: str, author: str, page: int, total: int, fmt: str) -> str:
    """Replace {title}, {author}, {page}, {totalPages} tokens."""
    pg_str    = _format_page_number(page, total, fmt)
    total_str = _format_page_number(total, total, fmt)
    return (
        text
        .replace("{title}",      doc_title)
        .replace("{author}",     author)
        .replace("{page}",       pg_str)
        .replace("{totalPages}", total_str)
    )


def _hf_cells(settings: dict, doc_title: str, author: str, page: int, total: int,
              left_field: str, center_field: str, right_field: str) -> tuple[str, str, str]:
    fmt = settings.get("pageNumberFormat", "number")
    left   = _resolve_hf_tokens(settings.get(left_field,   ""), doc_title, author, page, total, fmt)
    center = _resolve_hf_tokens(settings.get(center_field, ""), doc_title, author, page, total, fmt)
    right  = _resolve_hf_tokens(settings.get(right_field,  ""), doc_title, author, page, total, fmt)
    return left, center, right


def _page_num_cell(settings: dict, page: int, total: int, side: str) -> str:
    """Return the formatted page number if it belongs on `side` for this page."""
    if not settings.get("showPageNumbers", False):
        return ""
    pos = settings.get("pageNumberPosition", "center")
    # "none" was a legacy option — treat as disabled
    if pos in ("none", "", None):
        return ""
    if pos == "alternating":
        effective = "right" if page % 2 != 0 else "left"
    else:
        effective = pos
    if effective != side:
        return ""
    return _format_page_number(page, total, settings.get("pageNumberFormat", "number"))


# ── Chapter / scene title helpers ─────────────────────────────────────────────

_SPECIAL_SECTION_TITLES = {"prologue", "epilogue", "acknowledgements"}


def _is_special_chapter(chapter: dict) -> bool:
    return chapter.get("title", "").strip().lower() in _SPECIAL_SECTION_TITLES


def _format_scene_title(scene_title: str, settings: dict) -> str:
    """Return the scene title text, or empty string if scene titles are disabled."""
    if not settings.get("showSceneTitles", False):
        return ""
    return scene_title.strip()


def _format_chapter_title(order: int, title: str, settings: dict) -> str:
    fmt = settings.get("chapterTitleFormat", "none")
    if fmt == "none":
        return ""

    # Special sections: NO chapter number
    if title.strip().lower() in ("prologue", "epilogue", "acknowledgements"):
        return title.strip()

    # Normal chapters
    if fmt == "chapter-number":
        return f"Chapter {order}"
    if fmt == "chapter-number-title":
        return f"Chapter {order}: {title}"
    if fmt == "number-title":
        return f"{order}. {title}"
    if fmt == "title-only":
        return title

    return title

# ── DOCX export ────────────────────────────────────────────────────────────────

def build_docx(document_title: str, chapters: list, settings: dict) -> bytes:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    ALIGN_MAP = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }

    doc = DocxDocument()

    # ── Page size & margins ──────────────────────────────────────────
    section = doc.sections[0]
    paper = settings.get("paperFormat", "A4")
    if paper == "A4":
        section.page_width  = int(Cm(21.0))
        section.page_height = int(Cm(29.7))
    elif paper == "Letter":
        section.page_width  = int(Cm(21.59))
        section.page_height = int(Cm(27.94))
    elif paper == "A5":
        section.page_width  = int(Cm(14.8))
        section.page_height = int(Cm(21.0))
    elif paper == "Legal":
        section.page_width  = int(Cm(21.59))
        section.page_height = int(Cm(35.56))
    elif paper == "Custom":
        pw = settings.get("customWidth", 210)
        ph = settings.get("customHeight", 297)
        section.page_width  = int(Cm(pw / 10))
        section.page_height = int(Cm(ph / 10))

    unit = settings.get("marginUnit", "cm")
    section.top_margin    = Cm(_to_cm(settings.get("marginTop",    2.5), unit))
    section.bottom_margin = Cm(_to_cm(settings.get("marginBottom", 2.5), unit))
    section.left_margin   = Cm(_to_cm(settings.get("marginLeft",   2.5), unit))
    section.right_margin  = Cm(_to_cm(settings.get("marginRight",  2.5), unit))

    default_font_name = settings.get("defaultFont", "Arial, sans-serif").split(",")[0].strip().strip("'\"")
    doc_title_str = document_title or ""
    author_str    = settings.get("author", "") or ""
    default_font_pt   = float(settings.get("defaultFontSize", 12))
    default_align     = ALIGN_MAP.get(settings.get("defaultAlignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
    default_indent_cm = _parse_cm(
        f"{settings.get('defaultFirstLineIndent', 0)}{settings.get('defaultFirstLineIndentUnit', 'cm')}"
    )

    def _set_paragraph_spacing(para, line_height: float, paragraph_spacing_pt: float = None):
        """Set line spacing and paragraph spacing (space after) on a paragraph via XML."""
        pPr = para._p.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), str(int(line_height * 240)))
        spacing.set(qn("w:lineRule"), "auto")
        
        # Set space after (paragraph spacing) in twips (1/20th of a point)
        if paragraph_spacing_pt is not None:
            space_after_twips = int(paragraph_spacing_pt * 20)
            spacing.set(qn("w:after"), str(space_after_twips))
        
        pPr.append(spacing)

    def _apply_run_formatting(run, styles: dict, is_heading: bool = False, heading_pt: float = None):
        """Apply inline CSS styles to a docx Run."""
        fs_str = styles.get("font-size", "")
        if fs_str:
            run.font.size = Pt(_parse_pt(fs_str))
        elif heading_pt:
            run.font.size = Pt(heading_pt)
        else:
            run.font.size = Pt(default_font_pt)

        ff = styles.get("font-family", "")
        run.font.name = ff.split(",")[0].strip().strip("'\"") if ff else default_font_name

        fw = styles.get("font-weight", "")
        run.font.bold = fw in ("bold", "700", "800", "900") or is_heading

        fi = styles.get("font-style", "")
        run.font.italic = fi == "italic"

        td = styles.get("text-decoration", "")
        run.font.underline = "underline" in td
        run.font.strike    = "line-through" in td

    def _process_element(para, element, inherited_styles: dict = None, is_heading=False, heading_pt=None):
        """Recursively process a BeautifulSoup element and add runs to para."""
        inherited_styles = inherited_styles or {}

        if isinstance(element, NavigableString):
            text = str(element)
            if text:
                run = para.add_run(text)
                _apply_run_formatting(run, inherited_styles, is_heading, heading_pt)
            return

        tag = element.name or ""
        el_styles = dict(inherited_styles)
        el_styles.update(_parse_style(element.get("style", "")))

        # Handle bold/italic tags
        if tag in ("strong", "b"): el_styles["font-weight"] = "bold"
        if tag in ("em", "i"):     el_styles["font-style"]  = "italic"
        if tag in ("u",):          el_styles["text-decoration"] = "underline"
        if tag in ("s", "del"):    el_styles["text-decoration"] = "line-through"

        if tag == "br":
            run = para.add_run()
            run.add_break()
            return

        for child in element.children:
            _process_element(para, child, el_styles, is_heading, heading_pt)

    def _add_block(soup_el, first_block: bool = False):
        """Convert a block-level BeautifulSoup element to a docx paragraph."""
        tag = soup_el.name or "p"
        block_styles = _parse_style(soup_el.get("style", ""))

        # Heading detection
        heading_pt_map = {"h1": 28.0, "h2": 22.0, "h3": 18.0, "h4": 16.0}
        is_heading = tag in heading_pt_map
        heading_pt = heading_pt_map.get(tag)

        para = doc.add_paragraph()

        # Alignment
        align_str = block_styles.get("text-align", settings.get("defaultAlignment", "left"))
        para.alignment = ALIGN_MAP.get(align_str, default_align)

        # First-line indent
        ti_str = block_styles.get("text-indent", "")
        if ti_str and "var(--default-first-line-indent" not in ti_str:
            ti_cm = _parse_cm(ti_str)
            if ti_cm:
                para.paragraph_format.first_line_indent = Cm(ti_cm)
        elif default_indent_cm and not first_block:
            para.paragraph_format.first_line_indent = Cm(default_indent_cm)

        # Line height and paragraph spacing
        lh_str = block_styles.get("line-height", "")
        mb_str = block_styles.get("margin-bottom", "")
        try:
            lh = float(lh_str) if lh_str else float(settings.get("defaultLineHeight", 1.15))
            # Extract paragraph spacing from margin-bottom in pt
            ps_pt = None
            if mb_str and "pt" in mb_str:
                ps_pt = float(mb_str.replace("pt", "").strip())
            else:
                ps_pt = float(settings.get("defaultParagraphSpacing", 8))
            _set_paragraph_spacing(para, lh, ps_pt)
        except (ValueError, AttributeError):
            _set_paragraph_spacing(para, 1.15, 8)

        # Process inline content
        for child in soup_el.children:
            _process_element(para, child, {}, is_heading, heading_pt)

        return para

    # ── Title page or inline document title ─────────────────────────
    include_title_page = settings.get("includeTitlePage", False)
    tp_align_str = settings.get("titlePageAlignment", "center")
    tp_align     = ALIGN_MAP.get(tp_align_str, WD_ALIGN_PARAGRAPH.CENTER)

    def _tp_black_run(para, text: str, font_pt: float, bold: bool = False):
        """Add run(s) to para with black colour; splits on \\n to insert line breaks."""
        from docx.oxml import OxmlElement as _OxmlElement
        from docx.oxml.ns import qn as _qn
        lines = text.split("\n")
        for idx, line in enumerate(lines):
            if idx > 0:
                # Insert a soft line break (<w:br/>) between lines
                br_run = para.add_run()
                br_run.font.size = Pt(font_pt)
                br_run.font.bold = bold
                br_run.font.name = default_font_name
                br_run.font.color.rgb = RGBColor(0, 0, 0)
                br_elem = _OxmlElement("w:br")
                br_run._r.append(br_elem)
            run = para.add_run(line)
            run.font.size  = Pt(font_pt)
            run.font.bold  = bold
            run.font.name  = default_font_name
            run.font.color.rgb = RGBColor(0, 0, 0)

    def _tp_hf_para(cells: tuple[str, str, str], font_pt: float):
        """Render a three-cell (left/center/right) header or footer paragraph."""
        left_text, center_text, right_text = cells
        combined = " | ".join(t for t in (left_text, center_text, right_text) if t)
        if not combined:
            return
        # Approximate 3-column layout: left-aligned summary paragraph (DOCX
        # tab-stop approach matches what running H/F uses).
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.tab_stops.add_tab_stop(section.page_width // 2, WD_ALIGN_PARAGRAPH.CENTER)
        pf.tab_stops.add_tab_stop(section.page_width - section.right_margin, WD_ALIGN_PARAGRAPH.RIGHT)

        def _add_cell(text: str, add_tab: bool):
            if add_tab:
                t_run = para.add_run("\t")
                t_run.font.size = Pt(font_pt)
                t_run.font.name = default_font_name
            if text:
                _tp_black_run(para, text, font_pt)

        _add_cell(left_text,   add_tab=False)
        _add_cell(center_text, add_tab=True)
        _add_cell(right_text,  add_tab=True)

    if include_title_page:
        tp_h_pt = float(settings.get("titlePageHeaderFontSize", 10))
        tp_f_pt = float(settings.get("titlePageFooterFontSize", 10))
        tp_t_pt = float(settings.get("titlePageTitleFontSize",  28))
        tp_a_pt = float(settings.get("titlePageAuthorFontSize", 16))

        tp_h_cells = (
            (settings.get("titlePageHeaderLeft")   or "").strip(),
            (settings.get("titlePageHeaderCenter") or "").strip(),
            (settings.get("titlePageHeaderRight")  or "").strip(),
        )
        tp_f_cells = (
            (settings.get("titlePageFooterLeft")   or "").strip(),
            (settings.get("titlePageFooterCenter") or "").strip(),
            (settings.get("titlePageFooterRight")  or "").strip(),
        )

        # Header row
        _tp_hf_para(tp_h_cells, tp_h_pt)

        # Vertical spacer (approx 1/3 of page)
        for _ in range(8):
            doc.add_paragraph()

        # Title
        tp_title_text = (settings.get("titlePageTitle") or "").strip() or document_title
        tp_t_para = doc.add_paragraph()
        tp_t_para.alignment = tp_align
        _tp_black_run(tp_t_para, tp_title_text, tp_t_pt, bold=True)

        # Author — prefixed with "By "
        tp_author_raw  = (settings.get("titlePageAuthor") or "").strip() or author_str
        tp_author_text = f"By {tp_author_raw}" if tp_author_raw else ""
        if tp_author_text:
            doc.add_paragraph()
            tp_a_para = doc.add_paragraph()
            tp_a_para.alignment = tp_align
            _tp_black_run(tp_a_para, tp_author_text, tp_a_pt)

        # Vertical spacer
        for _ in range(8):
            doc.add_paragraph()

        # Footer row
        _tp_hf_para(tp_f_cells, tp_f_pt)

        doc.add_page_break()
    else:
        # Fallback: simple inline title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(document_title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.name = default_font_name
        doc.add_paragraph()  # blank line after title

    # ── Chapters ────────────────────────────────────────────────────
    for ch_idx, chapter in enumerate(chapters):
        ch_num = chapter.get("_chapter_number", ch_idx + 1)
        ch_title_text = _format_chapter_title(
            ch_num,
            chapter.get("title", ""),
            settings
        )

        if ch_title_text:
            title_style = settings.get("chapterTitleStyle", "bold")
            title_pt    = float(settings.get("chapterTitleSize", 16))
            title_align = ALIGN_MAP.get(settings.get("chapterTitleAlignment", "center"), WD_ALIGN_PARAGRAPH.CENTER)

            ct_para = doc.add_paragraph()
            ct_para.alignment = title_align
            ct_run = ct_para.add_run(ch_title_text)
            ct_run.font.size  = Pt(title_pt)
            ct_run.font.bold  = "bold" in title_style
            ct_run.font.italic = "italic" in title_style
            ct_run.font.name  = default_font_name

            is_special = _is_special_chapter(chapter)
            blank_lines = int(settings.get(
                "blankLinesAfterSpecialChapter" if is_special else "blankLinesAfterChapter", 2
            ))
            for _ in range(blank_lines):
                doc.add_paragraph()

        # ── Scenes ────────────────────────────────────────────────
        show_scene_titles      = settings.get("showSceneTitles", False) and not _is_special_chapter(chapter)
        scene_title_pt         = float(settings.get("sceneTitleSize", 13))
        scene_title_style      = settings.get("sceneTitleStyle", "bold")
        scene_title_align      = ALIGN_MAP.get(settings.get("sceneTitleAlignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
        blank_after_scene_title = int(settings.get("blankLinesAfterSceneTitle", 0))
        page_break_scene       = settings.get("pageBreakAfterSceneTitle", False)

        non_empty_scenes = [s for s in chapter.get("scenes", []) if s.get("content", "").strip()]
        for sc_idx, scene in enumerate(non_empty_scenes):
            content = scene.get("content", "")

            if show_scene_titles:
                if sc_idx > 0 and page_break_scene:
                    doc.add_page_break()
                sc_title_text = _format_scene_title(scene.get("title", ""), settings)
                if sc_title_text:
                    st_para = doc.add_paragraph()
                    st_para.alignment = scene_title_align
                    st_run = st_para.add_run(sc_title_text)
                    st_run.font.size   = Pt(scene_title_pt)
                    st_run.font.bold   = "bold" in scene_title_style
                    st_run.font.italic = "italic" in scene_title_style
                    st_run.font.name   = default_font_name
                    for _ in range(blank_after_scene_title):
                        doc.add_paragraph()
            elif sc_idx > 0:
                sep_para = doc.add_paragraph("* * *")
                sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Strip spacer divs before parsing
            content = re.sub(r'<div[^>]*data-type="page-break-spacer"[^>]*>.*?</div>', "", content, flags=re.DOTALL)

            soup = BeautifulSoup(content, "html.parser")
            block_tags = {"p", "h1", "h2", "h3", "h4", "h5", "h6", "blockquote", "div"}

            blocks = [el for el in soup.children if isinstance(el, Tag) and el.name in block_tags]
            if not blocks:
                if soup.get_text().strip():
                    para = doc.add_paragraph()
                    para.alignment = default_align
                    para.add_run(soup.get_text())
            else:
                for i, block in enumerate(blocks):
                    _add_block(block, first_block=(i == 0 and sc_idx == 0 and ch_idx == 0))

        # Page break between chapters
        if settings.get("pageBreakAfterChapter", True) and ch_idx < len(chapters) - 1:
            doc.add_page_break()

    # ── DOCX header / footer ────────────────────────────────────────
    # python-docx does not expose a page-count field easily, so we
    # use the {PAGE} / {NUMPAGES} DOCX field codes instead.
    def _add_hf_para(parent, left_text: str, center_text: str, right_text: str, font_pt: float):
        """Add a three-tab paragraph (left / center / right) to a header or footer."""
        from docx.oxml import OxmlElement as _OxmlElement
        from docx.oxml.ns import qn as _qn

        para = parent.paragraphs[0] if parent.paragraphs else parent.add_paragraph()
        para.clear()
        pf = para.paragraph_format
        pf.tab_stops.add_tab_stop(section.page_width // 2, WD_ALIGN_PARAGRAPH.CENTER)
        pf.tab_stops.add_tab_stop(section.page_width - section.right_margin, WD_ALIGN_PARAGRAPH.RIGHT)

        def _add_field(para, field_code: str):
            fld_char_begin = _OxmlElement("w:fldChar")
            fld_char_begin.set(_qn("w:fldCharType"), "begin")
            instr = _OxmlElement("w:instrText")
            instr.text = field_code
            fld_char_sep = _OxmlElement("w:fldChar")
            fld_char_sep.set(_qn("w:fldCharType"), "separate")
            fld_char_end = _OxmlElement("w:fldChar")
            fld_char_end.set(_qn("w:fldCharType"), "end")
            run = para.add_run()
            run._r.append(fld_char_begin)
            run._r.append(instr)
            run._r.append(fld_char_sep)
            run._r.append(fld_char_end)
            run.font.size = Pt(font_pt)
            run.font.name = default_font_name

        def _add_text_run(para, text: str):
            from docx.oxml import OxmlElement as _OxmlElement
            lines = text.split("\n")
            for idx, line in enumerate(lines):
                if idx > 0:
                    br_run = para.add_run()
                    br_run.font.size = Pt(font_pt)
                    br_run.font.name = default_font_name
                    br_run._r.append(_OxmlElement("w:br"))
                run = para.add_run(line)
                run.font.size = Pt(font_pt)
                run.font.name = default_font_name

        def _add_cell(para, text: str, add_tab: bool):
            if add_tab:
                run = para.add_run("\t")
                run.font.size = Pt(font_pt)
                run.font.name = default_font_name
            if text == "__PAGE__":
                _add_field(para, " PAGE ")
            elif text == "__NUMPAGES__":
                _add_field(para, " NUMPAGES ")
            elif "__PAGE_OF__" in text:
                _add_field(para, " PAGE ")
                _add_text_run(para, " of ")
                _add_field(para, " NUMPAGES ")
            elif text:
                _add_text_run(para, text)

        _add_cell(para, left_text,   add_tab=False)
        _add_cell(para, center_text, add_tab=True)
        _add_cell(para, right_text,  add_tab=True)

    def _encode_hf_text(raw: str, page_format: str) -> str:
        """Convert resolved token text back to DOCX field sentinels."""
        # Tokens are already resolved with dummy page numbers at this point
        # so we use the raw setting string instead.
        return raw

    def _build_hf_cells_docx(left_raw: str, center_raw: str, right_raw: str,
                              page_side_fn) -> tuple[str, str, str]:
        fmt  = settings.get("pageNumberFormat", "number")
        pg   = settings.get("pageNumberStart", 1)
        tot  = max(len(chapters), 1)

        def _resolve(text: str) -> str:
            text = text.replace("{title}", doc_title_str).replace("{author}", author_str)
            if "{page}" in text:
                text = text.replace("{page}", "__PAGE__")
            if "{totalPages}" in text:
                text = text.replace("{totalPages}", "__NUMPAGES__")
            return text

        return _resolve(left_raw), _resolve(center_raw), _resolve(right_raw)

    show_header = settings.get("showHeader", False)
    show_footer = settings.get("showFooter", False)
    show_pgnum  = settings.get("showPageNumbers", False)
    pgnum_pos   = settings.get("pageNumberPosition", "center")

    # When a title page is present, suppress H/F on the first physical page.
    if include_title_page and (show_header or show_footer or show_pgnum):
        section.different_first_page_header_footer = True

    if show_header or (show_pgnum and pgnum_pos != "none"):
        section.header_distance = Cm(0.5)
        if not include_title_page:
            section.different_first_page_header_footer = False
        header = section.header
        header.is_linked_to_previous = False

        h_left_raw   = settings.get("headerLeft",   "")
        h_center_raw = settings.get("headerCenter", "{title}")
        h_right_raw  = settings.get("headerRight",  "")
        h_font_pt    = float(settings.get("headerFontSize", 10))

        if not show_header:
            # Page numbers only — inject into the appropriate cell
            h_left_raw = h_center_raw = h_right_raw = ""
            if pgnum_pos == "left":       h_left_raw   = "{page}"
            elif pgnum_pos == "center":   h_center_raw = "{page}"
            elif pgnum_pos == "right":    h_right_raw  = "{page}"
            elif pgnum_pos == "alternating":
                # alternating requires separate odd/even — approximate with center
                h_center_raw = "{page}"
            h_font_pt = float(settings.get("pageNumberFontSize", 10))

        h_left, h_center, h_right = _build_hf_cells_docx(h_left_raw, h_center_raw, h_right_raw, None)
        _add_hf_para(header, h_left, h_center, h_right, h_font_pt)

    if show_footer or (show_pgnum and not show_header and pgnum_pos != "none"):
        section.footer_distance = Cm(0.5)
        footer = section.footer
        footer.is_linked_to_previous = False

        f_left_raw   = settings.get("footerLeft",   "")
        f_center_raw = settings.get("footerCenter", "")
        f_right_raw  = settings.get("footerRight",  "")
        f_font_pt    = float(settings.get("footerFontSize", 10))

        if not show_footer:
            f_left_raw = f_center_raw = f_right_raw = ""
            if pgnum_pos == "left":       f_left_raw   = "{page}"
            elif pgnum_pos == "center":   f_center_raw = "{page}"
            elif pgnum_pos == "right":    f_right_raw  = "{page}"
            elif pgnum_pos == "alternating":
                f_center_raw = "{page}"
            f_font_pt = float(settings.get("pageNumberFontSize", 10))

        f_left, f_center, f_right = _build_hf_cells_docx(f_left_raw, f_center_raw, f_right_raw, None)
        _add_hf_para(footer, f_left, f_center, f_right, f_font_pt)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF export ─────────────────────────────────────────────────────────────────
def build_pdf(document_title: str, chapters: list, settings: dict) -> bytes:
    from reportlab.lib.pagesizes import A4, LETTER, A5, LEGAL
    from reportlab.lib.units import cm, inch, mm
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from io import BytesIO
    import html as html_lib

    buf = BytesIO()

    # ── Page size ──────────────────────────────────────────────────
    paper = settings.get("paperFormat", "A4")
    if paper == "Letter":
        pagesize = LETTER
    elif paper == "A5":
        pagesize = A5
    elif paper == "Legal":
        pagesize = LEGAL
    elif paper == "Custom":
        pw = settings.get("customWidth", 210)
        ph = settings.get("customHeight", 297)
        pagesize = (pw * mm, ph * mm)
    else:
        pagesize = A4

    # ── Margins ────────────────────────────────────────────────────
    unit_str = settings.get("marginUnit", "cm")
    def to_points(val: float) -> float:
        if unit_str == "mm": return val * mm
        if unit_str == "in": return val * inch
        return val * cm

    margin_top    = to_points(settings.get("marginTop",    2.5))
    margin_bottom = to_points(settings.get("marginBottom", 2.5))
    margin_left   = to_points(settings.get("marginLeft",   2.5))
    margin_right  = to_points(settings.get("marginRight",  2.5))

    doc = SimpleDocTemplate(
        buf,
        pagesize=pagesize,
        topMargin=margin_top,
        bottomMargin=margin_bottom,
        leftMargin=margin_left,
        rightMargin=margin_right,
    )

    # ── Alignment ──────────────────────────────────────────────────
    ALIGN_MAP = {
        "left": TA_LEFT,
        "center": TA_CENTER,
        "right": TA_RIGHT,
        "justify": TA_JUSTIFY,
    }
    default_align    = ALIGN_MAP.get(settings.get("defaultAlignment", "left"), TA_LEFT)
    default_font_pt  = float(settings.get("defaultFontSize", 12))
    default_lh       = float(settings.get("defaultLineHeight", 1.15))
    default_indent   = _parse_cm(
        f"{settings.get('defaultFirstLineIndent', 0)}"
        f"{settings.get('defaultFirstLineIndentUnit', 'cm')}"
    ) * cm

    # ── Font name (reportlab uses Helvetica/Times-Roman/Courier) ──
    raw_font = settings.get("defaultFont", "Arial, sans-serif").split(",")[0].strip().strip("'\"").lower()
    if "times" in raw_font or "georgia" in raw_font or "serif" in raw_font:
        rl_font      = "Times-Roman"
        rl_font_bold = "Times-Bold"
        rl_font_it   = "Times-Italic"
        rl_font_bi   = "Times-BoldItalic"
    elif "courier" in raw_font or "mono" in raw_font:
        rl_font      = "Courier"
        rl_font_bold = "Courier-Bold"
        rl_font_it   = "Courier-Oblique"
        rl_font_bi   = "Courier-BoldOblique"
    else:
        rl_font      = "Helvetica"
        rl_font_bold = "Helvetica-Bold"
        rl_font_it   = "Helvetica-Oblique"
        rl_font_bi   = "Helvetica-BoldOblique"

    # ── Base paragraph style ───────────────────────────────────────
    default_para_spacing_pt = float(settings.get("defaultParagraphSpacing", 8))
    base_style = ParagraphStyle(
        "base",
        fontName=rl_font,
        fontSize=default_font_pt,
        leading=default_font_pt * default_lh * 1.2,
        alignment=default_align,
        firstLineIndent=default_indent,
        spaceAfter=default_para_spacing_pt,
        spaceBefore=0,
    )

    def _make_style(name, **kwargs) -> ParagraphStyle:
        s = ParagraphStyle(name, parent=base_style)
        for k, v in kwargs.items():
            setattr(s, k, v)
        return s

    # ── HTML → ReportLab XML converter ────────────────────────────
    # ReportLab's Paragraph accepts a limited subset of HTML-like tags.
    # We convert our stored HTML into that subset.

    def _html_to_rl(html_content: str, style: ParagraphStyle) -> list:
        """
        Parse HTML content and return a list of ReportLab flowables.
        Handles: p, h1-h4, br, strong, b, em, i, u, s, span (with inline styles).
        """
        # Strip spacer divs
        html_content = re.sub(
            r'<div[^>]*data-type="page-break-spacer"[^>]*>.*?</div>',
            "", html_content, flags=re.DOTALL
        )

        soup = BeautifulSoup(html_content, "html.parser")
        flowables = []

        heading_sizes = {"h1": 28, "h2": 22, "h3": 18, "h4": 16}

        def _node_to_rl_xml(node) -> str:
            """Convert a soup node to ReportLab-compatible XML string."""
            if isinstance(node, NavigableString):
                return html_lib.escape(str(node))

            tag = node.name or ""
            inner = "".join(_node_to_rl_xml(c) for c in node.children)

            if tag in ("strong", "b"):
                return f"<b>{inner}</b>"
            if tag in ("em", "i"):
                return f"<i>{inner}</i>"
            if tag == "u":
                return f"<u>{inner}</u>"
            if tag in ("s", "del"):
                return f"<strike>{inner}</strike>"
            if tag == "br":
                return "<br/>"
            if tag == "span":
                st = _parse_style(node.get("style", ""))
                parts = []
                fs = st.get("font-size", "")
                ff = st.get("font-family", "")
                fw = st.get("font-weight", "")
                fi = st.get("font-style", "")

                if fs:
                    try:
                        parts.append(f'size="{int(_parse_pt(fs))}"')
                    except Exception:
                        pass
                if ff:
                    fname = ff.split(",")[0].strip().strip("'\"").lower()
                    if "times" in fname or "georgia" in fname:
                        parts.append(f'face="Times-Roman"')
                    elif "courier" in fname or "mono" in fname:
                        parts.append(f'face="Courier"')
                    else:
                        parts.append(f'face="Helvetica"')

                wrapped = inner
                if fw in ("bold", "700", "800", "900"):
                    wrapped = f"<b>{wrapped}</b>"
                if fi == "italic":
                    wrapped = f"<i>{wrapped}</i>"

                if parts:
                    attrs = " ".join(parts)
                    return f'<font {attrs}>{wrapped}</font>'
                return wrapped

            # For block-level tags inside inline context, just return inner
            return inner

        block_tags = {"p", "h1", "h2", "h3", "h4", "h5", "h6", "blockquote"}

        for el in soup.children:
            if isinstance(el, NavigableString):
                text = str(el).strip()
                if text:
                    flowables.append(Paragraph(html_lib.escape(text), style))
                continue

            if not isinstance(el, Tag):
                continue

            tag = el.name or ""

            if tag in heading_sizes:
                hpt = heading_sizes[tag]
                h_style = _make_style(
                    f"h{tag}",
                    fontName=rl_font_bold,
                    fontSize=hpt,
                    leading=hpt * 1.3,
                    alignment=TA_LEFT,
                    firstLineIndent=0,
                    spaceBefore=hpt * 0.5,
                    spaceAfter=hpt * 0.3,
                )
                xml = "".join(_node_to_rl_xml(c) for c in el.children)
                flowables.append(Paragraph(xml or " ", h_style))
                continue

            if tag in block_tags or tag == "div":
                # Parse per-block styles
                st = _parse_style(el.get("style", ""))
                bl_align = ALIGN_MAP.get(st.get("text-align", ""), default_align)
                ti_str = st.get("text-indent", "")
                if ti_str and "var(" not in ti_str:
                    bl_indent = _parse_cm(ti_str) * cm
                else:
                    bl_indent = default_indent

                lh_str = st.get("line-height", "")
                try:
                    bl_lh = float(lh_str) if lh_str else default_lh
                except ValueError:
                    bl_lh = default_lh

                fs_str = st.get("font-size", "")
                bl_pt = _parse_pt(fs_str) if fs_str else default_font_pt

                p_style = _make_style(
                    f"p_{id(el)}",
                    alignment=bl_align,
                    firstLineIndent=bl_indent,
                    fontSize=bl_pt,
                    leading=bl_pt * bl_lh * 1.2,
                )

                xml = "".join(_node_to_rl_xml(c) for c in el.children)
                if xml.strip():
                    flowables.append(Paragraph(xml, p_style))
                else:
                    flowables.append(Spacer(1, default_font_pt))
                continue

        return flowables

    # ── Build story ────────────────────────────────────────────────
    story = []

    include_title_page_pdf = settings.get("includeTitlePage", False)
    if include_title_page_pdf:
        from reportlab.lib.colors import black as rl_black

        tp_align_str_pdf = settings.get("titlePageAlignment", "center")
        tp_rl_align      = ALIGN_MAP.get(tp_align_str_pdf, TA_CENTER)
        tp_title_text    = (settings.get("titlePageTitle") or "").strip() or document_title
        tp_author_raw    = (settings.get("titlePageAuthor") or "").strip() or (settings.get("author") or "")
        tp_author_text   = f"By {tp_author_raw}" if tp_author_raw else ""
        tp_title_pt      = float(settings.get("titlePageTitleFontSize",  28))
        tp_author_pt     = float(settings.get("titlePageAuthorFontSize", 16))
        tp_header_pt     = float(settings.get("titlePageHeaderFontSize", 10))
        tp_footer_pt     = float(settings.get("titlePageFooterFontSize", 10))

        tp_hL = (settings.get("titlePageHeaderLeft")   or "").strip()
        tp_hC = (settings.get("titlePageHeaderCenter") or "").strip()
        tp_hR = (settings.get("titlePageHeaderRight")  or "").strip()
        tp_fL = (settings.get("titlePageFooterLeft")   or "").strip()
        tp_fC = (settings.get("titlePageFooterCenter") or "").strip()
        tp_fR = (settings.get("titlePageFooterRight")  or "").strip()

        # Black text style factory for title page
        def _tp_style(name, fn, pt, align):
            s = _make_style(name, fontName=fn, fontSize=pt, leading=pt * 1.4,
                            alignment=align, firstLineIndent=0, spaceAfter=0)
            s.textColor = rl_black
            return s

        # Three-column H/F line: approximate with a tab-stop paragraph via XML
        # For simplicity we render as "left | center | right" plain text joined
        def _tp_hf_line(l: str, c: str, r: str, pt: float):
            parts = [x for x in (l, c, r) if x]
            if not parts:
                return
            # Convert \n to <br/> for ReportLab XML Paragraphs
            def _rl_escape(s: str) -> str:
                return html_lib.escape(s).replace("\n", "<br/>")
            combined = "   ".join(_rl_escape(p) for p in parts)
            st = _tp_style(f"tp_hf_{id((l,c,r))}", rl_font, pt, tp_rl_align)
            story.append(Paragraph(combined, st))

        ph = doc.pagesize[1]

        _tp_hf_line(tp_hL, tp_hC, tp_hR, tp_header_pt)

        spacer_h = (ph - doc.topMargin - doc.bottomMargin) * 0.30
        story.append(Spacer(1, max(spacer_h, 12)))

        story.append(Paragraph(
            html_lib.escape(tp_title_text),
            _tp_style("tp_title", rl_font_bold, tp_title_pt, tp_rl_align),
        ))
        if tp_author_text:
            story.append(Spacer(1, tp_author_pt * 0.5))
            story.append(Paragraph(
                html_lib.escape(tp_author_text),
                _tp_style("tp_author", rl_font, tp_author_pt, tp_rl_align),
            ))

        spacer_h2 = (ph - doc.topMargin - doc.bottomMargin) * 0.30
        story.append(Spacer(1, max(spacer_h2, 12)))

        _tp_hf_line(tp_fL, tp_fC, tp_fR, tp_footer_pt)

        story.append(PageBreak())
    else:
        # Simple inline document title
        title_style = _make_style(
            "doc_title",
            fontName=rl_font_bold,
            fontSize=24,
            leading=30,
            alignment=TA_CENTER,
            firstLineIndent=0,
            spaceAfter=24,
        )
        story.append(Paragraph(html_lib.escape(document_title), title_style))
        story.append(Spacer(1, 12))

    ch_title_style_raw = settings.get("chapterTitleStyle", "bold")
    ch_title_pt        = float(settings.get("chapterTitleSize", 16))
    ch_title_align     = ALIGN_MAP.get(settings.get("chapterTitleAlignment", "center"), TA_CENTER)
    ch_title_fn        = rl_font_bold if "bold" in ch_title_style_raw else rl_font
    if "italic" in ch_title_style_raw:
        ch_title_fn = rl_font_bi if "bold" in ch_title_style_raw else rl_font_it

    ch_title_style = _make_style(
        "ch_title",
        fontName=ch_title_fn,
        fontSize=ch_title_pt,
        leading=ch_title_pt * 1.4,
        alignment=ch_title_align,
        firstLineIndent=0,
        spaceBefore=ch_title_pt,
        spaceAfter=ch_title_pt * 0.5,
    )

    show_scene_titles     = settings.get("showSceneTitles", False)
    scene_title_pt        = float(settings.get("sceneTitleSize", 13))
    sc_title_style_val    = settings.get("sceneTitleStyle", "bold")
    sc_title_align_val    = settings.get("sceneTitleAlignment", "left")
    sc_title_rl_align     = ALIGN_MAP.get(sc_title_align_val, TA_LEFT)
    sc_title_fn           = rl_font_bold if "bold" in sc_title_style_val else rl_font
    if "italic" in sc_title_style_val:
        sc_title_fn = rl_font_bi if "bold" in sc_title_style_val else rl_font_it
    sc_title_rl_style = _make_style(
        "sc_title",
        fontName=sc_title_fn,
        fontSize=scene_title_pt,
        leading=scene_title_pt * 1.4,
        alignment=sc_title_rl_align,
        firstLineIndent=0,
        spaceBefore=scene_title_pt * 0.5,
        spaceAfter=scene_title_pt * 0.25,
    )
    scene_sep_style = _make_style(
        "scene_sep",
        fontName=rl_font,
        fontSize=default_font_pt,
        leading=default_font_pt * default_lh,
        alignment=TA_CENTER,
        firstLineIndent=0,
        spaceBefore=default_font_pt * default_lh,
        spaceAfter=default_font_pt * default_lh,
    )

    for ch_idx, chapter in enumerate(chapters):
        ch_num = chapter.get("_chapter_number", ch_idx + 1)
        ch_title_text = _format_chapter_title(
            ch_num,
            chapter.get("title", ""),
            settings
        )

        if ch_title_text:
            story.append(Paragraph(html_lib.escape(ch_title_text), ch_title_style))
            is_special_ch = _is_special_chapter(chapter)
            blank_lines = int(settings.get(
                "blankLinesAfterSpecialChapter" if is_special_ch else "blankLinesAfterChapter", 2
            ))
            for _ in range(blank_lines):
                story.append(Spacer(1, default_font_pt * default_lh))

        chapter_show_sc_titles  = show_scene_titles and not _is_special_chapter(chapter)
        blank_after_sc_title_pdf = int(settings.get("blankLinesAfterSceneTitle", 0))
        page_break_scene_pdf     = settings.get("pageBreakAfterSceneTitle", False)
        non_empty_scenes = [s for s in chapter.get("scenes", []) if s.get("content", "").strip()]
        for sc_idx, scene in enumerate(non_empty_scenes):
            content = scene.get("content", "")
            if chapter_show_sc_titles:
                if sc_idx > 0 and page_break_scene_pdf:
                    story.append(PageBreak())
                sc_title_text = _format_scene_title(scene.get("title", ""), settings)
                if sc_title_text:
                    story.append(Paragraph(html_lib.escape(sc_title_text), sc_title_rl_style))
                    for _ in range(blank_after_sc_title_pdf):
                        story.append(Spacer(1, default_font_pt * default_lh))
            elif sc_idx > 0:
                story.append(Paragraph("* * *", scene_sep_style))
            flowables = _html_to_rl(content, base_style)
            story.extend(flowables)

        if settings.get("pageBreakAfterChapter", True) and ch_idx < len(chapters) - 1:
            story.append(PageBreak())

    # ── PDF header / footer via canvas callbacks ───────────────────
    show_header_pdf = settings.get("showHeader", False)
    show_footer_pdf = settings.get("showFooter", False)
    show_pgnum_pdf  = settings.get("showPageNumbers", False)
    pgnum_pos_pdf   = settings.get("pageNumberPosition", "center")
    pgnum_fmt_pdf   = settings.get("pageNumberFormat", "number")
    pgnum_start_pdf = int(settings.get("pageNumberStart", 1))
    pg_font_pt_pdf  = float(settings.get("pageNumberFontSize", 10))
    h_font_pt_pdf   = float(settings.get("headerFontSize", 10))
    f_font_pt_pdf   = float(settings.get("footerFontSize", 10))
    doc_title_pdf   = document_title or ""
    author_pdf      = settings.get("author", "") or ""

    total_pages_pdf: list[int] = [0]  # filled on second pass

    def _draw_hf(canvas_obj, doc_obj):
        canvas_obj.saveState()
        page_idx  = doc_obj.page - 1
        # When a title page is present it occupies physical page 1; skip H/F there
        # and offset the display page numbers for the rest.
        if include_title_page_pdf and doc_obj.page == 1:
            canvas_obj.restoreState()
            return
        body_offset = 1 if include_title_page_pdf else 0
        disp_page = (page_idx - body_offset) + pgnum_start_pdf
        total     = total_pages_pdf[0] or (doc_obj.page - body_offset)
        fmt       = pgnum_fmt_pdf
        pw, ph    = canvas_obj._pagesize

        def _pg_num_for_side(side: str) -> str:
            if not show_pgnum_pdf or pgnum_pos_pdf == "none":
                return ""
            if pgnum_pos_pdf == "alternating":
                eff = "right" if disp_page % 2 != 0 else "left"
            else:
                eff = pgnum_pos_pdf
            if eff != side:
                return ""
            return _format_page_number(disp_page, total, fmt)

        def _resolve(text: str) -> str:
            return _resolve_hf_tokens(text, doc_title_pdf, author_pdf, disp_page, total, fmt)

        def _draw_band(y: float, left_raw: str, center_raw: str, right_raw: str,
                       font_pt: float, pg_side_override: bool):
            left   = _resolve(left_raw)   or (pg_side_override and _pg_num_for_side("left")   or "")
            center = _resolve(center_raw) or (pg_side_override and _pg_num_for_side("center") or "")
            right  = _resolve(right_raw)  or (pg_side_override and _pg_num_for_side("right")  or "")

            canvas_obj.setFont(rl_font, font_pt)
            canvas_obj.setFillColorRGB(0.4, 0.4, 0.4)
            line_gap = font_pt * 1.4

            def _draw_multiline_left(text: str, base_y: float):
                for i, line in enumerate(text.split("\n")):
                    canvas_obj.drawString(margin_left, base_y - i * line_gap, line)

            def _draw_multiline_center(text: str, base_y: float):
                for i, line in enumerate(text.split("\n")):
                    canvas_obj.drawCentredString(pw / 2, base_y - i * line_gap, line)

            def _draw_multiline_right(text: str, base_y: float):
                for i, line in enumerate(text.split("\n")):
                    canvas_obj.drawRightString(pw - margin_right, base_y - i * line_gap, line)

            if left:
                _draw_multiline_left(left, y)
            if center:
                _draw_multiline_center(center, y)
            if right:
                _draw_multiline_right(right, y)

        if show_header_pdf or (show_pgnum_pdf and pgnum_pos_pdf != "none"):
            h_y = ph - margin_top + h_font_pt_pdf * 0.8
            if show_header_pdf:
                _draw_band(h_y,
                           settings.get("headerLeft", ""),
                           settings.get("headerCenter", "{title}"),
                           settings.get("headerRight", ""),
                           h_font_pt_pdf, show_pgnum_pdf)
            else:
                _draw_band(h_y, "", "", "", pg_font_pt_pdf, True)

        if show_footer_pdf or (show_pgnum_pdf and not show_header_pdf and pgnum_pos_pdf != "none"):
            f_y = margin_bottom - f_font_pt_pdf * 1.5
            if show_footer_pdf:
                _draw_band(f_y,
                           settings.get("footerLeft", ""),
                           settings.get("footerCenter", ""),
                           settings.get("footerRight", ""),
                           f_font_pt_pdf, show_pgnum_pdf and not show_header_pdf)
            else:
                _draw_band(f_y, "", "", "", pg_font_pt_pdf, True)

        canvas_obj.restoreState()

    def _on_first_page(canvas_obj, doc_obj):
        _draw_hf(canvas_obj, doc_obj)

    def _on_later_pages(canvas_obj, doc_obj):
        _draw_hf(canvas_obj, doc_obj)

    doc.build(story, onFirstPage=_on_first_page, onLaterPages=_on_later_pages)

    # Second pass: embed actual total-page count if {totalPages} was used
    # (ReportLab doesn't support forward-references natively in SimpleDocTemplate,
    # so we accept the approximation; total_pages_pdf[0] stays 0 meaning we
    # fall back to doc.page which is the last page number — good enough.)
    return buf.getvalue()


def build_epub(document_title: str, chapters: list, settings: dict) -> bytes:
    import logging
    import html as _html

    log = logging.getLogger(__name__)

    try:
        from ebooklib import epub
    except ImportError as exc:
        raise RuntimeError(
            "EPUB export requires ebooklib. Install with `pip install ebooklib`."
        ) from exc

    book = epub.EpubBook()
    book.set_identifier(str(uuid.uuid4()))
    book.set_title(document_title)
    language = settings.get("language") or "en"
    book.set_language(language)

    author = settings.get("author") or settings.get("creator")
    if author:
        book.add_author(str(author))

    default_font = settings.get("defaultFont", "Arial, sans-serif").split(",")[0].strip().strip("'\"")
    default_font_size = float(settings.get("defaultFontSize", 12))
    default_line_height = float(settings.get("defaultLineHeight", 1.15))
    default_indent_cm = _parse_cm(
        f"{settings.get('defaultFirstLineIndent', 0)}{settings.get('defaultFirstLineIndentUnit', 'cm')}"
    )

    margin_unit = settings.get("marginUnit", "cm")
    margin_top = _to_cm(settings.get("marginTop", 2.5), margin_unit)
    margin_bottom = _to_cm(settings.get("marginBottom", 2.5), margin_unit)
    margin_left = _to_cm(settings.get("marginLeft", 2.5), margin_unit)
    margin_right = _to_cm(settings.get("marginRight", 2.5), margin_unit)

    chapter_title_size = float(settings.get("chapterTitleSize", 16))
    chapter_title_alignment = settings.get("chapterTitleAlignment", "center")
    chapter_title_style = settings.get("chapterTitleStyle", "bold")
    chapter_title_weight = "bold" if "bold" in str(chapter_title_style) else "normal"
    chapter_title_italic = "italic" if "italic" in str(chapter_title_style) else "normal"
    chapter_blank_lines = int(settings.get("blankLinesAfterChapter", 2))
    blank_margin = max(1, chapter_blank_lines) * 0.8

    scene_title_size_epub   = float(settings.get("sceneTitleSize", 13))
    scene_title_align_epub  = settings.get("sceneTitleAlignment", "left")
    scene_title_style_epub  = settings.get("sceneTitleStyle", "bold")
    sc_title_weight         = "bold" if "bold" in str(scene_title_style_epub) else "normal"
    sc_title_italic         = "italic" if "italic" in str(scene_title_style_epub) else "normal"
    sc_blank_after          = int(settings.get("blankLinesAfterSceneTitle", 0))
    sc_blank_margin         = max(0.25, sc_blank_after) * 0.8 if sc_blank_after > 0 else 0.25

    # ── EPUB header/footer settings ────────────────────────────────
    show_header_epub  = settings.get("showHeader", False)
    show_footer_epub  = settings.get("showFooter", False)
    show_pgnum_epub   = settings.get("showPageNumbers", False)
    pgnum_pos_epub    = settings.get("pageNumberPosition", "center")
    h_font_pt_epub    = float(settings.get("headerFontSize", 10))
    f_font_pt_epub    = float(settings.get("footerFontSize", 10))
    pg_font_pt_epub   = float(settings.get("pageNumberFontSize", 10))

    # EPUB @page counter-based page numbers (CSS Paged Media)
    page_counter_css = ""
    if show_pgnum_epub and pgnum_pos_epub != "none":
        align_map = {"left": "left", "center": "center", "right": "right",
                     "alternating": "center", "none": "center"}
        pg_align = align_map.get(pgnum_pos_epub, "center")
        page_counter_css = f"""
    @page {{
      @bottom-{pg_align} {{
        content: counter(page);
        font-size: {pg_font_pt_epub}pt;
        font-family: '{default_font}', serif;
        color: #555;
      }}
    }}"""

    css = f"""
    body {{
      font-family: '{default_font}', serif;
      font-size: {default_font_size}pt;
      line-height: {default_line_height};
      margin: 0;
      padding: {margin_top}cm {margin_right}cm {margin_bottom}cm {margin_left}cm;
    }}{page_counter_css}
    p {{
      text-indent: {default_indent_cm}cm;
      margin: 0 0 1em 0;
    }}
    .chapter-title {{
      font-size: {chapter_title_size}pt;
      font-weight: {chapter_title_weight};
      font-style: {chapter_title_italic};
      text-align: {chapter_title_alignment};
      margin: 0 0 {blank_margin}em 0;
    }}
    .scene-title {{
      font-size: {scene_title_size_epub}pt;
      font-weight: {sc_title_weight};
      font-style: {sc_title_italic};
      text-align: {scene_title_align_epub};
      margin: 0.75em 0 {sc_blank_margin}em 0;
      text-indent: 0;
    }}
    .scene-page-break {{
      page-break-before: always;
      break-before: page;
      height: 0;
      margin: 0;
      padding: 0;
    }}
    .scene-sep {{
      text-align: center;
      margin: 1em 0;
      text-indent: 0;
    }}
    .doc-header {{
      font-size: {h_font_pt_epub}pt;
      color: #555;
      border-bottom: 1px solid #ccc;
      display: flex;
      justify-content: space-between;
      padding-bottom: 0.25em;
      margin-bottom: 1em;
      text-indent: 0;
    }}
    .doc-footer {{
      font-size: {f_font_pt_epub}pt;
      color: #555;
      border-top: 1px solid #ccc;
      display: flex;
      justify-content: space-between;
      padding-top: 0.25em;
      margin-top: 1em;
      text-indent: 0;
    }}
    .scene {{
      margin-bottom: 1.25em;
    }}
    h1, h2, h3, h4, h5, h6 {{
      margin: 0.8em 0 0.4em 0;
    }}
    """

    style_item = epub.EpubItem(
        uid="style",
        file_name="styles/epub_styles.css",
        media_type="text/css",
        content=css.encode("utf-8"),
    )
    book.add_item(style_item)

    def _make_xhtml(chapter_label: str, body_html: str) -> bytes:
        """
        Wrap body content in a complete, well-formed XHTML document and return
        UTF-8 bytes. lxml requires bytes (not str) when an XML encoding
        declaration is present; passing a str causes "Document is empty".
        The stylesheet link is embedded here because ebooklib skips its template
        when page.content is pre-set.
        """
        doc = (
            "<?xml version='1.0' encoding='utf-8'?>"
            "<html xmlns='http://www.w3.org/1999/xhtml'>"
            "<head>"
            "<meta charset='utf-8'/>"
            f"<title>{_html.escape(chapter_label)}</title>"
            "<link rel='stylesheet' type='text/css' href='../styles/epub_styles.css'/>"
            "</head>"
            f"<body>{body_html}</body>"
            "</html>"
        )
        return doc.encode("utf-8")

    spine = ["nav"]
    toc = []
    pages_created = 0

    # ── EPUB title page ──────────────────────────────────────────────
    include_title_page_epub = settings.get("includeTitlePage", False)
    if include_title_page_epub:
        tp_align_epub   = settings.get("titlePageAlignment", "center")
        tp_title_epub   = _html.escape((settings.get("titlePageTitle") or "").strip() or document_title)
        tp_author_raw_e = (settings.get("titlePageAuthor") or "").strip() or (settings.get("author") or "")
        tp_author_epub  = _html.escape(f"By {tp_author_raw_e}" if tp_author_raw_e else "")
        tp_title_pt_e   = float(settings.get("titlePageTitleFontSize",  28))
        tp_author_pt_e  = float(settings.get("titlePageAuthorFontSize", 16))
        tp_header_pt_e  = float(settings.get("titlePageHeaderFontSize", 10))
        tp_footer_pt_e  = float(settings.get("titlePageFooterFontSize", 10))

        tp_hL_e = (settings.get("titlePageHeaderLeft")   or "").strip()
        tp_hC_e = (settings.get("titlePageHeaderCenter") or "").strip()
        tp_hR_e = (settings.get("titlePageHeaderRight")  or "").strip()
        tp_fL_e = (settings.get("titlePageFooterLeft")   or "").strip()
        tp_fC_e = (settings.get("titlePageFooterCenter") or "").strip()
        tp_fR_e = (settings.get("titlePageFooterRight")  or "").strip()

        has_tp_header = any([tp_hL_e, tp_hC_e, tp_hR_e])
        has_tp_footer = any([tp_fL_e, tp_fC_e, tp_fR_e])

        # Three-column band matching the running H/F layout
        def _epub_nl2br(s: str) -> str:
            import html as _html_e
            return _html_e.escape(s).replace("\n", "<br/>")

        def _epub_tp_band(l: str, c: str, r: str, pt: float) -> str:
            return (
                f"<div style='display:flex;justify-content:space-between;"
                f"font-size:{pt}pt;color:#000000;line-height:1.4;'>"
                f"<span style='flex:1;text-align:left;white-space:pre-wrap;'>{_epub_nl2br(l)}</span>"
                f"<span style='flex:1;text-align:center;white-space:pre-wrap;'>{_epub_nl2br(c)}</span>"
                f"<span style='flex:1;text-align:right;white-space:pre-wrap;'>{_epub_nl2br(r)}</span>"
                f"</div>"
            )

        tp_parts = []
        if has_tp_header:
            tp_parts.append(_epub_tp_band(tp_hL_e, tp_hC_e, tp_hR_e, tp_header_pt_e))

        tp_parts.append(
            f"<div style='flex:1;display:flex;flex-direction:column;justify-content:center;"
            f"text-align:{tp_align_epub};'>"
            f"<p style='font-size:{tp_title_pt_e}pt;font-weight:bold;color:#000000;"
            f"margin:0 0 0.5em 0;text-indent:0;'>{tp_title_epub}</p>"
        )
        if tp_author_epub:
            tp_parts.append(
                f"<p style='font-size:{tp_author_pt_e}pt;color:#000000;"
                f"margin:0;text-indent:0;'>{tp_author_epub}</p>"
            )
        tp_parts.append("</div>")

        if has_tp_footer:
            tp_parts.append(_epub_tp_band(tp_fL_e, tp_fC_e, tp_fR_e, tp_footer_pt_e))

        tp_body_html = (
            "<div style='min-height:90vh;display:flex;flex-direction:column;"
            "color:#000000;font-size:inherit;'>"
            + "\n".join(tp_parts)
            + "</div>"
        )
        tp_page = epub.EpubHtml(
            title="Title Page",
            file_name="title_page.xhtml",
            lang=language,
        )
        tp_page.content = _make_xhtml("Title Page", tp_body_html)
        book.add_item(tp_page)
        toc.append(epub.Link(tp_page.file_name, "Title Page", tp_page.file_name))
        spine.append(tp_page)
        pages_created += 1

    for ch_idx, chapter in enumerate(chapters):
        chapter_title = chapter.get("title", "") or f"Chapter {ch_idx + 1}"
        ch_num = chapter.get("_chapter_number", ch_idx + 1)
        formatted_title = _format_chapter_title(ch_num, chapter_title, settings)
        chapter_label = formatted_title or chapter_title

        chapter_body = []
        if formatted_title:
            chapter_body.append(
                f"<h1 class='chapter-title'>{_html.escape(formatted_title)}</h1>"
            )

        show_sc_titles          = settings.get("showSceneTitles", False) and not _is_special_chapter(chapter)
        page_break_scene_epub   = settings.get("pageBreakAfterSceneTitle", False)
        has_scene_content = False
        sc_idx = 0
        for scene in chapter.get("scenes", []):
            raw_content = scene.get("content", "") or ""
            clean = _sanitize_html_to_xhtml(raw_content)

            # _sanitize_html_to_xhtml returns the &#160; placeholder only when
            # the raw content was truly empty — skip those scenes.
            if not raw_content.strip():
                log.debug(
                    "EPUB: chapter %d '%s' — skipping empty scene '%s'",
                    ch_idx + 1, chapter_title, scene.get("title", ""),
                )
                continue

            if show_sc_titles:
                if sc_idx > 0 and page_break_scene_epub:
                    chapter_body.append("<div class='scene-page-break'></div>")
                sc_title_text = _format_scene_title(scene.get("title", ""), settings)
                if sc_title_text:
                    chapter_body.append(
                        f"<p class='scene-title'>{_html.escape(sc_title_text)}</p>"
                    )
            elif sc_idx > 0:
                chapter_body.append("<p class='scene-sep'>* * *</p>")

            chapter_body.append(f"<div class='scene'>{clean}</div>")
            has_scene_content = True
            sc_idx += 1

        if not has_scene_content and not formatted_title:
            log.warning(
                "EPUB: chapter %d '%s' has no content and no title — skipping",
                ch_idx + 1, chapter_title,
            )
            continue

        body_html = "\n".join(chapter_body)
        if not body_html.strip():
            body_html = "<p>&#160;</p>"

        # Inject header/footer bands (EPUB doesn't support true per-page headers,
        # but we add them once per chapter as a visual approximation).
        hf_doc_title = document_title or ""
        hf_author    = settings.get("author", "") or ""
        hf_fmt       = settings.get("pageNumberFormat", "number")
        hf_pg_start  = int(settings.get("pageNumberStart", 1))
        hf_total     = max(pages_created + 1, 1)
        hf_disp_page = ch_idx + hf_pg_start

        def _epub_hf_cell(text: str) -> str:
            resolved = _resolve_hf_tokens(text, hf_doc_title, hf_author, hf_disp_page, hf_total, hf_fmt)
            escaped = _html.escape(resolved).replace("\n", "<br/>")
            return f"<span style='white-space:pre-wrap;'>{escaped}</span>"

        if show_header_epub:
            h_l = settings.get("headerLeft",   "")
            h_c = settings.get("headerCenter", "{title}")
            h_r = settings.get("headerRight",  "")
            header_band = (
                f"<div class='doc-header'>"
                f"{_epub_hf_cell(h_l)}{_epub_hf_cell(h_c)}{_epub_hf_cell(h_r)}"
                f"</div>"
            )
            body_html = header_band + "\n" + body_html

        if show_footer_epub:
            f_l = settings.get("footerLeft",   "")
            f_c = settings.get("footerCenter", "")
            f_r = settings.get("footerRight",  "")
            footer_band = (
                f"<div class='doc-footer'>"
                f"{_epub_hf_cell(f_l)}{_epub_hf_cell(f_c)}{_epub_hf_cell(f_r)}"
                f"</div>"
            )
            body_html = body_html + "\n" + footer_band

        page = epub.EpubHtml(
            title=chapter_label,
            file_name=f"chapter_{ch_idx + 1}.xhtml",
            lang=language,
        )
        page.content = _make_xhtml(chapter_label, body_html)

        book.add_item(page)
        toc.append(epub.Link(page.file_name, chapter_label, page.file_name))
        spine.append(page)
        pages_created += 1

    if pages_created == 0:
        log.warning(
            "EPUB: document '%s' has no exportable content — inserting placeholder",
            document_title,
        )
        placeholder = epub.EpubHtml(
            title="Empty Document",
            file_name="chapter_1.xhtml",
            lang=language,
        )
        placeholder.content = _make_xhtml(
            "Empty Document", "<p>This document has no content.</p>"
        )
        book.add_item(placeholder)
        toc.append(epub.Link(placeholder.file_name, "Empty Document", placeholder.file_name))
        spine.append(placeholder)

    book.toc = toc
    book.spine = spine
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    buf = io.BytesIO()
    epub.write_epub(buf, book)
    return buf.getvalue()